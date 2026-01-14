from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:color'), '2C5F8D')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_subheader(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(33, 37, 41)

def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(10)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F5F5F5')
    p._element.get_or_add_pPr().append(shading)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '16')
    left.set(qn('w:color'), '2C5F8D')
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(33,37,41)

def create_sp_doc(spName, version, createdDate, createdBy, schema, objectType, purpose, 
                  recentChanges, whatsNew, parameters, logicFlow, dependencies, 
                  usageExamples, performanceNotes, errorHandling, fullVersionHistory, 
                  complexityScore, isQA=False):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Header
    header_table = doc.add_table(rows=1, cols=2)
    for row in header_table.rows:
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border in ['top','left','bottom','right','insideH','insideV']:
                b = OxmlElement(f'w:{border}')
                b.set(qn('w:val'), 'none')
                tcBorders.append(b)
            tcPr.append(tcBorders)
    
    left_cell = header_table.rows[0].cells[0]
    left_cell.width = Inches(4.5)
    p = left_cell.paragraphs[0]
    
    doc_title = "QA STORED PROCEDURE" if isQA else "STORED PROCEDURE"
    r = p.add_run(doc_title)
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = RGBColor(44,95,141)
    
    p2 = left_cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    r2 = p2.add_run("Technical Documentation")
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(73,80,87)
    
    p3 = left_cell.add_paragraph()
    p3.paragraph_format.space_before = Pt(6)
    r3 = p3.add_run(f"{schema}.{spName}")
    r3.font.size = Pt(10)
    r3.font.bold = True
    r3.font.color.rgb = RGBColor(73,80,87)
    
    right_cell = header_table.rows[0].cells[1]
    right_cell.width = Inches(2.5)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F8F9FA')
    right_cell._element.get_or_add_tcPr().append(shading)
    tcPr = right_cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border in ['top','left','bottom','right']:
        b = OxmlElement(f'w:{border}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '8')
        b.set(qn('w:color'), 'DEE2E6')
        tcBorders.append(b)
    tcPr.append(tcBorders)
    right_cell.paragraphs[0].text = ""
    
    qa_label = "QA Procedure" if isQA else "Production"
    metadata = [
        ("Version:", f"v{version}", RGBColor(73,80,87)),
        ("Type:", qa_label, RGBColor(40,167,69) if isQA else RGBColor(73,80,87)),
        ("Created:", createdDate, RGBColor(73,80,87)),
        ("Created By:", createdBy, RGBColor(73,80,87)),
        ("Complexity:", f"{complexityScore}/100", RGBColor(73,80,87))
    ]
    for label, value, color in metadata:
        p = right_cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(label + " ")
        r1.font.size = Pt(9)
        r1.font.bold = True
        r1.font.color.rgb = RGBColor(73,80,87)
        r2 = p.add_run(value)
        r2.font.size = Pt(9)
        r2.font.bold = True
        r2.font.color.rgb = color
    
    add_divider(doc)
    
    # Recent Changes (5 most recent)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("5 MOST RECENT CHANGES")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor(44,95,141)
    
    if recentChanges:
        for change in recentChanges[:5]:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Inches(0.25)
            run = p.add_run(f"{change['date']} - {change['summary']} ({change['refDoc']})")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(73,80,87)
    else:
        p = doc.add_paragraph("No changes recorded yet.")
        p.paragraph_format.space_after = Pt(4)
        run = p.runs[0]
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(108,117,125)
    
    add_divider(doc)
    
    # Section 1: Purpose
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("1. PURPOSE")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor(44,95,141)
    
    if isQA:
        qa_purpose_intro = "This is a QA validation procedure designed to verify data quality and integrity. "
        p = doc.add_paragraph(qa_purpose_intro + purpose)
    else:
        p = doc.add_paragraph(purpose)
    p.paragraph_format.space_after = Pt(10)
    p.runs[0].font.size = Pt(11)
    
    add_divider(doc)
    
    # Section 2: What's New (only if recent version)
    if whatsNew:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(f"2. WHAT'S NEW IN VERSION {version}")
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = RGBColor(44,95,141)
        
        add_subheader(doc, "Changes in This Version:")
        p = doc.add_paragraph(whatsNew)
        p.paragraph_format.space_after = Pt(10)
        p.runs[0].font.size = Pt(11)
        
        add_divider(doc)
        section_num = 3
    else:
        section_num = 2
    
    # Section: Parameters
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(f"{section_num}. PARAMETERS")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor(44,95,141)
    
    if parameters:
        for param in parameters:
            add_subheader(doc, f"{param['name']} ({param['type']}):")
            p = doc.add_paragraph(f"{param['description']}")
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.left_indent = Inches(0.25)
            p.runs[0].font.size = Pt(10)
            if 'defaultValue' in param and param['defaultValue']:
                p2 = doc.add_paragraph(f"Default: {param['defaultValue']}")
                p2.paragraph_format.space_after = Pt(8)
                p2.paragraph_format.left_indent = Inches(0.25)
                p2.runs[0].font.size = Pt(9)
                p2.runs[0].font.color.rgb = RGBColor(108,117,125)
    else:
        p = doc.add_paragraph("No parameters")
        p.runs[0].font.size = Pt(10)
    
    add_divider(doc)
    section_num += 1
    
    # Section: Logic Flow
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(f"{section_num}. LOGIC FLOW")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor(44,95,141)
    
    if isinstance(logicFlow, list):
        for i, step in enumerate(logicFlow, 1):
            add_subheader(doc, f"Step {i}: {step['title']}")
            p = doc.add_paragraph(step['description'])
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.left_indent = Inches(0.25)
            p.runs[0].font.size = Pt(10)
    else:
        p = doc.add_paragraph(logicFlow)
        p.runs[0].font.size = Pt(11)
    
    add_divider(doc)
    section_num += 1
    
    # Section: Dependencies (only if complexity > 30)
    if complexityScore > 30 and dependencies:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(f"{section_num}. DEPENDENCIES")
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = RGBColor(44,95,141)
        
        if 'tables' in dependencies and dependencies['tables']:
            add_subheader(doc, "Tables Accessed:")
            for table in dependencies['tables']:
                p = doc.add_paragraph(f"• {table}", style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.25)
                p.runs[0].font.size = Pt(10)
        
        if 'procedures' in dependencies and dependencies['procedures']:
            add_subheader(doc, "Stored Procedures Called:")
            for proc in dependencies['procedures']:
                p = doc.add_paragraph(f"• {proc}", style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.25)
                p.runs[0].font.size = Pt(10)
        
        add_divider(doc)
        section_num += 1
    
    # Section: Usage Examples
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(f"{section_num}. USAGE EXAMPLES")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor(44,95,141)
    
    if usageExamples:
        for i, example in enumerate(usageExamples, 1):
            add_subheader(doc, f"Example {i}: {example['title']}")
            add_code_block(doc, example['code'])
            if 'explanation' in example:
                p = doc.add_paragraph(example['explanation'])
                p.paragraph_format.space_after = Pt(10)
                p.paragraph_format.left_indent = Inches(0.25)
                p.runs[0].font.size = Pt(10)
    
    add_divider(doc)
    section_num += 1
    
    # Section: Performance Notes (only if complexity > 50)
    if complexityScore > 50 and performanceNotes:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(f"{section_num}. PERFORMANCE NOTES")
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = RGBColor(44,95,141)
        
        p = doc.add_paragraph(performanceNotes)
        p.runs[0].font.size = Pt(11)
        
        add_divider(doc)
        section_num += 1
    
    # Section: Error Handling (only if complexity > 40)
    if complexityScore > 40 and errorHandling:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(f"{section_num}. ERROR HANDLING")
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = RGBColor(44,95,141)
        
        p = doc.add_paragraph(errorHandling)
        p.runs[0].font.size = Pt(11)
        
        add_divider(doc)
        section_num += 1
    
    # Section: Full Version History (at end)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(f"{section_num}. FULL VERSION HISTORY")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor(44,95,141)
    
    if fullVersionHistory:
        history_table = doc.add_table(rows=len(fullVersionHistory)+1, cols=4)
        history_table.style = 'Light Grid Accent 1'
        
        header_cells = history_table.rows[0].cells
        headers = ['Version', 'Date', 'Changed By', 'Changes']
        for i, header_text in enumerate(headers):
            header_cells[i].text = header_text
            for p in header_cells[i].paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
        
        for i, entry in enumerate(fullVersionHistory, 1):
            row = history_table.rows[i]
            row.cells[0].text = f"v{entry['version']}"
            row.cells[1].text = entry['date']
            row.cells[2].text = entry['changedBy']
            row.cells[3].text = f"{entry['changes']} (Ref: {entry['refDoc']})" if entry['refDoc'] else entry['changes']
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
    
    return doc

# Generate document
if __name__ == "__main__":
    sp_doc = create_sp_doc(
        spName="usp_Customer_Update",
        version="1.2",
        createdDate="2024-10-01",
        createdBy="A.Kirby",
        schema="dbo",
        objectType="Stored Procedure",
        purpose="Updates customer information including contact details and preferences. Validates input data and maintains audit trail.",
        recentChanges=[
            {"date": "2024-12-03", "summary": "Added error handling for NULL inputs", "refDoc": "DF-0089"},
            {"date": "2024-11-15", "summary": "Optimized JOIN performance", "refDoc": "EN-0067"},
            {"date": "2024-10-20", "summary": "Added email validation", "refDoc": "BR-0045"},
            {"date": "2024-10-10", "summary": "Fixed timezone handling", "refDoc": "DF-0032"},
            {"date": "2024-10-01", "summary": "Initial documentation", "refDoc": ""}
        ],
        whatsNew="Added comprehensive error handling for NULL and invalid input parameters. Procedure now validates email format before update and returns detailed error codes.",
        parameters=[
            {"name": "@CustomerID", "type": "INT", "description": "Unique customer identifier", "defaultValue": None},
            {"name": "@Email", "type": "VARCHAR(255)", "description": "Customer email address", "defaultValue": None},
            {"name": "@Phone", "type": "VARCHAR(20)", "description": "Customer phone number", "defaultValue": "NULL"}
        ],
        logicFlow=[
            {"title": "Input Validation", "description": "Validates all input parameters. Checks CustomerID exists, email format is valid, phone number format is correct."},
            {"title": "Begin Transaction", "description": "Starts transaction to ensure data consistency across multiple table updates."},
            {"title": "Update Customer Table", "description": "Updates primary customer record with new contact information."},
            {"title": "Update Audit Log", "description": "Records change in audit table with timestamp and user information."},
            {"title": "Commit Transaction", "description": "Commits all changes if successful, rolls back on any error."}
        ],
        dependencies={
            "tables": ["dbo.Customers", "dbo.CustomerAudit", "dbo.EmailValidation"],
            "procedures": ["dbo.usp_ValidateEmail", "dbo.usp_LogAuditEntry"]
        },
        usageExamples=[
            {
                "title": "Update customer email",
                "code": "EXEC dbo.usp_Customer_Update \n    @CustomerID = 12345,\n    @Email = 'john.doe@example.com',\n    @Phone = NULL",
                "explanation": "Updates email address for customer ID 12345 without changing phone number."
            }
        ],
        performanceNotes="Procedure uses covering index on CustomerID for optimal lookup performance. Average execution time: 15ms. Recommended to batch updates if processing >1000 records.",
        errorHandling="Returns error code -1 for invalid CustomerID, -2 for invalid email format, -3 for database constraint violations. All errors are logged to ErrorLog table.",
        fullVersionHistory=[
            {"version": "1.2", "date": "2024-12-03", "changedBy": "A.Kirby", "changes": "Added error handling for NULL inputs", "refDoc": "DF-0089"},
            {"version": "1.1", "date": "2024-11-15", "changedBy": "J.Smith", "changes": "Optimized JOIN performance", "refDoc": "EN-0067"},
            {"version": "1.0", "date": "2024-10-01", "changedBy": "System", "changes": "Initial documentation", "refDoc": ""}
        ],
        complexityScore=45,
        isQA=False
    )
    
    sp_doc.save("TEMPLATE_StoredProcedure.docx")
    print("SP template created: TEMPLATE_StoredProcedure.docx")
